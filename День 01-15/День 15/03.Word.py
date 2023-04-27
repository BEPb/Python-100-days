"""
Python 3.10 создание Word документа
Название файла '03.Word.py'

Version: 0.1
Author: Andrej Marinchenko
Date: 2023-04-27
"""
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Оглавление документа', 0)

p = document.add_paragraph('Простой абзац, имеющий некоторые символы ')
p.add_run('жирные').bold = True
p.add_run(' а некоторые ')
p.add_run('курсивом.').italic = True

document.add_heading('Заголовок, уровень 1', level=1)
document.add_paragraph('Интенсивная цитата', style='Intense Quote')

document.add_paragraph(
    'первый элемент в неупорядоченном списке', style='List Bullet'
)
document.add_paragraph(
    'первый элемент в упорядоченном списке', style='List Number'
)

document.add_picture('./res/guido.jpg', width=Inches(1.25))

records = (
    (3, '101', 'Спам'),
    (7, '422', 'тест'),
    (4, '631', 'проверка')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'столбец 1'
hdr_cells[1].text = 'столбец 2'
hdr_cells[2].text = 'столбец 3'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('./res/demo.docx')