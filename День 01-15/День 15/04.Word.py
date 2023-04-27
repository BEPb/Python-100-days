"""
Python 3.9 чтение Word документа
Название файла '04.Word.py'

Version: 0.1
Author: Andrej Marinchenko
Date: 2023-04-27
"""

from docx import Document

doc = Document('./res/test.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))
